from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from sqlalchemy.orm import Session
from app.database import SessionLocal
from app.models.assemblee import Assemblea, PuntoOrdineGiorno, PartecipazioneAssemblea
from app.models.utenti import Tesserato
from app.schemas.assemblee import AssembleaCreate, AssembleaRead, PuntoOrdineGiornoCreate, PuntoOrdineGiornoRead, PartecipazioneCreate, PartecipazioneRead
from typing import List
import os
import cloudinary
import cloudinary.uploader

router = APIRouter(tags=["Assemblee"])

cloudinary.config(cloudinary_url=os.getenv("CLOUDINARY_URL"))

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# ---- ASSEMBLEE ----

@router.get("/assemblee/", response_model=List[AssembleaRead])
def lista_assemblee(db: Session = Depends(get_db)):
    return db.query(Assemblea).all()

@router.get("/assemblee/{assemblea_id}", response_model=AssembleaRead)
def get_assemblea(assemblea_id: int, db: Session = Depends(get_db)):
    assemblea = db.query(Assemblea).filter(Assemblea.id == assemblea_id).first()
    if not assemblea:
        raise HTTPException(status_code=404, detail="Assemblea non trovata")
    return assemblea

@router.post("/assemblee/", response_model=AssembleaRead)
def crea_assemblea(assemblea: AssembleaCreate, db: Session = Depends(get_db)):
    db_assemblea = Assemblea(**assemblea.model_dump())
    db.add(db_assemblea)
    db.commit()
    db.refresh(db_assemblea)
    return db_assemblea

@router.put("/assemblee/{assemblea_id}", response_model=AssembleaRead)
def aggiorna_assemblea(assemblea_id: int, assemblea: AssembleaCreate, db: Session = Depends(get_db)):
    db_assemblea = db.query(Assemblea).filter(Assemblea.id == assemblea_id).first()
    if not db_assemblea:
        raise HTTPException(status_code=404, detail="Assemblea non trovata")
    for key, value in assemblea.model_dump().items():
        setattr(db_assemblea, key, value)
    db.commit()
    db.refresh(db_assemblea)
    return db_assemblea

@router.delete("/assemblee/{assemblea_id}")
def elimina_assemblea(assemblea_id: int, db: Session = Depends(get_db)):
    db_assemblea = db.query(Assemblea).filter(Assemblea.id == assemblea_id).first()
    if not db_assemblea:
        raise HTTPException(status_code=404, detail="Assemblea non trovata")
    db.delete(db_assemblea)
    db.commit()
    return {"messaggio": "Assemblea eliminata"}

# ---- PUNTI ORDINE DEL GIORNO ----

@router.get("/assemblee/{assemblea_id}/punti", response_model=List[PuntoOrdineGiornoRead])
def lista_punti(assemblea_id: int, db: Session = Depends(get_db)):
    return db.query(PuntoOrdineGiorno).filter(PuntoOrdineGiorno.assemblea_id == assemblea_id).all()

@router.post("/punti/", response_model=PuntoOrdineGiornoRead)
def crea_punto(punto: PuntoOrdineGiornoCreate, db: Session = Depends(get_db)):
    db_punto = PuntoOrdineGiorno(**punto.model_dump())
    db.add(db_punto)
    db.commit()
    db.refresh(db_punto)
    return db_punto

@router.put("/punti/{punto_id}/esito", response_model=PuntoOrdineGiornoRead)
def aggiorna_esito(punto_id: int, esito: str, db: Session = Depends(get_db)):
    db_punto = db.query(PuntoOrdineGiorno).filter(PuntoOrdineGiorno.id == punto_id).first()
    if not db_punto:
        raise HTTPException(status_code=404, detail="Punto non trovato")
    db_punto.esito = esito
    db.commit()
    db.refresh(db_punto)
    return db_punto

# ---- PARTECIPAZIONI ----

@router.get("/assemblee/{assemblea_id}/partecipanti", response_model=List[PartecipazioneRead])
def lista_partecipanti(assemblea_id: int, db: Session = Depends(get_db)):
    return db.query(PartecipazioneAssemblea).filter(PartecipazioneAssemblea.assemblea_id == assemblea_id).all()

@router.post("/partecipazioni/", response_model=PartecipazioneRead)
def registra_partecipazione(partecipazione: PartecipazioneCreate, db: Session = Depends(get_db)):
    db_part = PartecipazioneAssemblea(**partecipazione.model_dump())
    db.add(db_part)
    db.commit()
    db.refresh(db_part)
    return db_part

@router.put("/partecipazioni/{partecipazione_id}", response_model=PartecipazioneRead)
def aggiorna_partecipazione(partecipazione_id: int, partecipazione: PartecipazioneCreate, db: Session = Depends(get_db)):
    db_part = db.query(PartecipazioneAssemblea).filter(PartecipazioneAssemblea.id == partecipazione_id).first()
    if not db_part:
        raise HTTPException(status_code=404, detail="Partecipazione non trovata")
    for key, value in partecipazione.model_dump().items():
        setattr(db_part, key, value)
    db.commit()
    db.refresh(db_part)
    return db_part


# ---- VERBALE ----

@router.post("/assemblee/{assemblea_id}/verbale", response_model=AssembleaRead)
async def carica_verbale(assemblea_id: int, file: UploadFile = File(...), db: Session = Depends(get_db)):
    db_assemblea = db.query(Assemblea).filter(Assemblea.id == assemblea_id).first()
    if not db_assemblea:
        raise HTTPException(status_code=404, detail="Assemblea non trovata")

    contenuto = await file.read()
    nome_base, estensione = os.path.splitext(file.filename)
    nome_base_pulito = nome_base.replace(' ', '_')
    public_id_finale = f"verbale_{assemblea_id}_{nome_base_pulito}{estensione}"
    risultato = cloudinary.uploader.upload(
        contenuto,
        folder="gestionale/assemblee/verbali",
        resource_type="raw",
        public_id=public_id_finale,
        use_filename=False,
        unique_filename=True,
        overwrite=True,
        type="upload",
        access_mode="public",
    )
    db_assemblea.path_verbale = risultato["secure_url"]
    db.commit()
    db.refresh(db_assemblea)
    return db_assemblea


@router.delete("/assemblee/{assemblea_id}/verbale", response_model=AssembleaRead)
def elimina_verbale(assemblea_id: int, db: Session = Depends(get_db)):
    db_assemblea = db.query(Assemblea).filter(Assemblea.id == assemblea_id).first()
    if not db_assemblea:
        raise HTTPException(status_code=404, detail="Assemblea non trovata")
    db_assemblea.path_verbale = None
    db.commit()
    db.refresh(db_assemblea)
    return db_assemblea


@router.get("/assemblee/{assemblea_id}/verbale/genera-docx")
def genera_verbale_docx(assemblea_id: int, db: Session = Depends(get_db)):
    """Genera una bozza di verbale in formato Word (.docx), precompilata con
    i dati dell'assemblea, i punti all'ordine del giorno e i partecipanti,
    pronta per essere completata e modificata liberamente prima di essere
    eventualmente caricata firmata."""
    from fastapi.responses import StreamingResponse
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    assemblea = db.query(Assemblea).filter(Assemblea.id == assemblea_id).first()
    if not assemblea:
        raise HTTPException(status_code=404, detail="Assemblea non trovata")
    punti = db.query(PuntoOrdineGiorno).filter(PuntoOrdineGiorno.assemblea_id == assemblea_id).order_by(PuntoOrdineGiorno.numero.asc()).all()
    partecipazioni = db.query(PartecipazioneAssemblea).filter(PartecipazioneAssemblea.assemblea_id == assemblea_id).all()

    doc = Document()
    for sezione in doc.sections:
        sezione.top_margin = Cm(2)
        sezione.bottom_margin = Cm(2)

    titolo = doc.add_paragraph()
    titolo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titolo.add_run("VERBALE DI ASSEMBLEA\nASD PGS JUVENILIA")
    run.bold = True
    run.font.size = Pt(15)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(f"Assemblea: ").bold = True
    p.add_run(assemblea.titolo)
    p = doc.add_paragraph()
    p.add_run("Data: ").bold = True
    p.add_run(assemblea.data.strftime("%d/%m/%Y") if assemblea.data else "____________")
    if assemblea.ora:
        p.add_run("   Ora: ").bold = True
        p.add_run(str(assemblea.ora))
    p = doc.add_paragraph()
    p.add_run("Luogo: ").bold = True
    p.add_run(assemblea.luogo or "____________________________")

    doc.add_paragraph()
    doc.add_heading("Partecipanti", level=2)
    if partecipazioni:
        tabella = doc.add_table(rows=1, cols=3)
        tabella.style = "Light Grid Accent 1"
        hdr = tabella.rows[0].cells
        hdr[0].text = "Cognome e Nome"
        hdr[1].text = "Presente"
        hdr[2].text = "Delega"
        for part in partecipazioni:
            tess = db.query(Tesserato).filter(Tesserato.id == part.tesserato_id).first()
            delegato = db.query(Tesserato).filter(Tesserato.id == part.delega_a_id).first() if part.delega_a_id else None
            riga = tabella.add_row().cells
            riga[0].text = f"{tess.cognome} {tess.nome}" if tess else "-"
            riga[1].text = "Sì" if part.presente else "No"
            riga[2].text = f"{delegato.cognome} {delegato.nome}" if delegato else "-"
    else:
        doc.add_paragraph("Nessun partecipante ancora registrato nel sistema per questa assemblea.")

    doc.add_paragraph()
    doc.add_heading("Ordine del giorno", level=2)
    if punti:
        for pt in punti:
            doc.add_paragraph(f"{pt.numero}. {pt.titolo}", style="List Number")
            if pt.descrizione:
                doc.add_paragraph(pt.descrizione)
            doc.add_paragraph("Discussione ed esito:").italic = True
            doc.add_paragraph(pt.esito or "________________________________________________")
            doc.add_paragraph()
    else:
        doc.add_paragraph("1. ________________________________________________")
        doc.add_paragraph("Discussione ed esito:")
        doc.add_paragraph("________________________________________________")

    doc.add_paragraph()
    doc.add_heading("Delibere", level=2)
    doc.add_paragraph("________________________________________________")
    doc.add_paragraph("________________________________________________")

    doc.add_paragraph()
    doc.add_paragraph()
    firme = doc.add_paragraph()
    firme.add_run("Il Presidente").bold = True
    firme.add_run("\t\t\t\t")
    firme.add_run("Il Segretario").bold = True
    doc.add_paragraph("_____________________")
    doc.add_paragraph()
    doc.add_paragraph("_____________________\t\t\t_____________________")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    nome_file = f"verbale_bozza_{assemblea.titolo.replace(' ', '_')}.docx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{nome_file}"'}
    )
